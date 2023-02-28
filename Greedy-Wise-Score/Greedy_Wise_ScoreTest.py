import matplotlib
import operator
import random
import scipy.stats as stats
import pandas as pd
import numpy
import copy
import json
import ast
from anytree import PreOrderIter
import math
import itertools
import scipy.stats as stats
import matplotlib.pyplot as pl
from operator import itemgetter, attrgetter, methodcaller
from scipy.stats import poisson
from docx import Document
from docx.text.run import Font, Run
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from time import time
from anytree import AnyNode
from anytree.exporter import JsonExporter
import Greedy_Wise_Utility
from numpy.random import rand
from numpy.random import seed
from scipy.stats import kendalltau
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import t



def timer_func(func):
    # This function shows the execution time of
    # the function object passed
    def wrap_func(*args, **kwargs):
        t1 = time()
        result = func(*args, **kwargs)
        t2 = time()
        print(f'Function {func.__name__!r} executed in {(t2-t1):.4f}s')
        return result
    return wrap_func



'''
The multinomial CDF function is the implementation of Levin's "Representaion of Multinomial Cumulative Distribution Function"

G:= Number of total groups (including non-protected)
k:= Position
p:= Array of probabilities of each group to be selected
tau_p:= Array of number of protected items

EXAMPLE:
Groups: ["White"(NP), "Black"(P), "Asian"(P), "Hispanic"(P)]
k: 30
p: [0.4, 0.3, 0.2, 0.1]
tau_p: [k, 10, 5, 3]
The ranking is fair if multinomCDF(4, 30, p, tau_p) > a = 0.1
'''

def multinomCDF_log(G, k, p, tau_p):
    s = float(k);
    log_cdf = -poisson.logpmf(k,s);
    gamma1 = 0.0;
    gamma2 = 0.0;
    sum_s2 = 0.0;
    sum_mu = 0.0;

    # P(W=k)
    for i in range(0,G):
        i = i-1
        sp = s*p[i];

        pcdf = poisson.cdf(tau_p[i],sp);
        log_cdf += numpy.log(pcdf);

        mu = sp*(1-poisson.pmf(tau_p[i],sp)/pcdf);
        s2 = mu-(tau_p[i]-mu)*(sp-mu);

        mr = tau_p[i];
        mf2 = sp*mu-mr*(sp-mu);

        mr *= tau_p[i]-1;
        mf3 = sp*mf2-mr*(sp-mu);

        mr *= tau_p[i]-2;
        mf4 = sp*mf3-mr*(sp-mu);

        mu2 = mf2+mu*(1-mu);
        mu3 = mf3+mf2*(3-3*mu)+mu*(1+mu*(-3+2*mu));
        mu4 = mf4+mf3*(6-4*mu)+mf2*(7+mu*(-12+6*mu))+mu*(1+mu*(-4+mu*(6-3*mu)));

        gamma1 += mu3;
        gamma2 += mu4-3*s2*s2;
        sum_mu += mu;
        sum_s2 += s2;
    sp = numpy.sqrt(sum_s2);
    gamma1 /= sum_s2*sp;
    gamma2 /= sum_s2*sum_s2;

    x = (k-sum_mu)/sp;
    x2 = x*x;

    PWN = (-x2/2
    +numpy.log(1+gamma1/6*x*(x2-3)+gamma2/24*(x2*x2-6*x2+3)
    +gamma1*gamma1/72*(((x2-15)*x2+45)*x2-15))
    -numpy.log(2*math.pi)/2 -numpy.log(sp));

    log_cdf += PWN;
    return log_cdf;

def multinomCDF(G, k, p, tau_p):
    return numpy.exp(multinomCDF_log(G, k, p, tau_p ));

"""
To find a continuous multinomial icdf solution from a solution space.
Used for the Ranking Algorithm.
If the multinomial_CDF with tau alreay results in the value > a, return tau.
Else calculate the multinomial_CDF of tau in each case when each element of tau is increased by 1.
Compare the value and take the tau that produces the closest multinomial_CDF to a, which is also > a.

For example, initial tau: [0, 0, 0, 0]
if:
    multinomial_CDF(G, k, p, tau) > a, return tau
else:
    compare:
        multinomial_CDF(G, k, p, [1, 0, 0, 0])
        multinomial_CDF(G, k, p, [0, 1, 0, 0])
        multinomial_CDF(G, k, p, [0, 0, 1, 0])
        multinomial_CDF(G, k, p, [0, 0, 0, 1])
    and take the increased tau with its multinomial_CDF > a and closest to a

"""
def multinomial_icdf_continuous(G, k, p, a, tau):
    tau_p = [k] + list(tau);
    temp = copy.copy(tau_p)
    cdf = multinomCDF(G, k, p, tau_p)
    new_cdf = 0;
    initial = 1;
    not_fulfilled = 0;

    if(cdf > a):
        return tau_p;
    for i in range(len(tau_p)-1):
        temp[i+1] = temp[i+1]+1;
        if(initial == 1):
            tau_p = copy.copy(temp);
            cdf = multinomCDF(G, k, p, tau_p);
            initial = 0;
        else:
            new_cdf = multinomCDF(G, k, p, temp)
            if(new_cdf >= a and new_cdf >= cdf):
                tau_p = copy.copy(temp);
                cdf = multinomCDF(G, k, p, tau_p);
        if(new_cdf >= a or cdf >= a):
            temp[i+1] = temp[i+1]-1
    return tau_p

"""
Finds the target group which does not achieve the minimum target.
"""
def find_target(minimum_targets, count, categories):
    for i in range(len(minimum_targets)):
        if(minimum_targets[i] > count[i+1]):
                return categories[i+1];

def find_achieved_target(minimum_targets, count, categories):
    for i in range(len(minimum_targets)):
        if(minimum_targets[i] < count[i+1]):
                return categories[i+1];

def get_num_categories(attributeNamesAndCategories):
    num_categories = 1
    for i in attributeNamesAndCategories.items():
        num_categories *= i[1]
    return num_categories

def determineGroups(attributeNamesAndCategories):
    elementSets = []
    groups = []
    for attr, cardinality in attributeNamesAndCategories.items():
        elementSets.append(list(range(0, cardinality)))

    groups = list(itertools.product(*elementSets))
    return groups
@timer_func
def get_minimum_targets(categories, p, alpha, k):
    positions = numpy.array(list(range(k))) + 1;
    minimum_targets = [];
    tau = numpy.zeros(len(categories)-1);

    for i in positions:
        tau_p = multinomial_icdf_continuous(len(p), i, p , alpha, tau)[1:]
        minimum_targets.append(numpy.array(tau_p));
        tau = copy.copy(tau_p);
    df = pd.DataFrame(data=(numpy.array(minimum_targets)).astype(int))
    df.columns = p[1:]
    df.index = numpy.array(range(k))+1
    df.to_html("minimum_target_table.html")
    return minimum_targets


@timer_func
def get_minimum_targets2(categories, p, alpha, k):
    positions = numpy.array(list(range(k))) + 1;
    minimum_targets = [];
    tau = numpy.zeros(len(categories)-1);
    block_sizes = []
    count=0

    for i in positions:
        tau_p = multinomial_icdf_continuous(len(p), i, p , alpha, tau)[1:]

        if(all(tau_p[j]==tau[j] for j in range(len(tau)))):
            count = count + 1
        else:
            block_sizes.append(count)
            count =1
        minimum_targets.append(numpy.array(tau_p));
        print ("tau_p: ",tau_p, " cdf: ",multinomCDF(len(categories),i, p, [i]+tau_p))
        tau = copy.copy(tau_p);

    block_sizes.append(count)
    return minimum_targets, block_sizes


def separate_groups(data_set, categories, attributeItems):
    num_categories = len(categories)
    separateByGroups = [[] for _ in range(num_categories)]

    for i in data_set:
        categorieList = []
        for j in attributeItems:
            categorieList.append(i[j[0]])
        separateByGroups[categories.index(tuple(categorieList))].append(i)
        categorieList = []
    return separateByGroups

def plot(data_set,attributeNamesAndCategories, attributeQuality):
    colors = ['black', 'red', 'blue', 'green', 'yellow','slategray', 'darkseagreen', 'indigo']
    markers = ['-o','-<','-s','-+', '-d', '-h', '-p', '-8']
#     label=['Germany','Turkey','Greece','Italy','Spain','France']
    best = data_set[0][attributeQuality]
    categories = determineGroups(attributeNamesAndCategories)
    attributeItems = attributeNamesAndCategories.items()
    output_ranking_separated = separate_groups(data_set,categories, attributeItems)
    separateQualityByGroups = []
    fig = pl.figure(figsize=(20, 10))
    pl.subplot(211)
    round_2f = []
    for k in data_set:
        k[attributeQuality] = float(k[attributeQuality])/best

    for i in range(len(categories)):
        separateQualityByGroups.append([quality[attributeQuality] for quality in output_ranking_separated[i]])
        fit = stats.norm.pdf(separateQualityByGroups[i], numpy.mean(separateQualityByGroups[i]), numpy.std(separateQualityByGroups[i]))
        pl.plot(separateQualityByGroups[i],fit,markers[i], markersize=6, label=categories[i], color=colors[i])
#         pl.plot(separateQualityByGroups[i],fit,markers[i], markersize=6, label=label[i], color=colors[i])
        pl.legend(loc='center left', fontsize='x-large', bbox_to_anchor=(1, 0.5))
        round_2f.append([round(elem, 2) for elem in separateQualityByGroups[i]])
    pl.xlabel(attributeQuality+' (Quality)')
    pl.ylabel('Probability Density Function')

    pl.subplot(212)
    pl.hist(round_2f,30, histtype='bar', label=categories, color=colors[:len(categories)])
#     pl.hist(round_2f,30, histtype='bar', label=label, color=colors[:len(categories)])
    pl.xlabel(attributeQuality+' (Quality)')
    pl.ylabel('Frequency')
    pl.legend(loc='center left', fontsize='x-large', bbox_to_anchor=(1, 0.5))
    pl.show()
    out_png = '../Plots/histo_plot.png'
    pl.savefig(out_png, dpi=100)

def plot_scatter(data_set,attributeNamesAndCategories, attributeQuality):
#     label=['Germany','Turkey','Greece','Italy','Spain','France']

    colors = ['black', 'red', 'blue', 'green', 'yellow', 'magenta','slategray', 'darkseagreen', 'indigo']
    markers = ['o','<','s','>', 'd', 'x', 'h', 'p', '8']
    categories = determineGroups(attributeNamesAndCategories)
    attributeItems = attributeNamesAndCategories.items()
    separateQualityByGroups = []
    positionByGroups =[]
    for i in range(len(data_set)):
        data_set[i]['k']=i+1
    output_ranking_separated = separate_groups(data_set,categories, attributeItems)
    fig = pl.figure(figsize=(6, 6), dpi=100)
    for i in range(len(categories)):
        separateQualityByGroups.append([quality[attributeQuality] for quality in output_ranking_separated[i]])
        positionByGroups.append([k['k'] for k in output_ranking_separated[i]])
        area = numpy.pi*4
        pl.scatter(positionByGroups[i], separateQualityByGroups[i], s=area, color=colors[i], alpha=0.5, label=categories[i], marker=markers[i])
#         pl.scatter(positionByGroups[i], separateQualityByGroups[i], s=area, color=colors[i], alpha=0.5, label=label[i], marker=markers[i])
    pl.xlabel('Position (k)')
    pl.ylabel(attributeQuality+' (Quality)')
    pl.legend(loc='center left', fontsize='x-large', bbox_to_anchor=(1, 0.5))
#     pl.ylim(0, 6.0)
    pl.show()
    out_png = '../Plots/scatter.png'
    pl.savefig(out_png, dpi=100)


"""
    - data_set: list of items in dictionary form, which specifies its group and quality
        ex)
        [{'Gender': 0, 'Uni score': 0.5344585435},
         {'Gender': 2, 'Uni score': 1.0160173543},
         {'Gender': 0, 'Uni score': 0.4573686375},
         {'Gender': 1, 'Uni score': 0.6513702354}]

    - attributeNamesAndCategories: attributes that are being used to specify groups in the data_set and its number of categories
        ex)
        {"Gender": 4}

    - attributeQuality: the name of the attribute used for quality
        ex)
        "Uni score"

    - k: total length of the ranking

    - p: array of prbability, that each group could be selected into the ranking
        ex)
        [0.4,0.3,0.2,0.1]
            -> 0.4 for Gender: 0 (Non-protected)
            -> 0.3 for Gender: 1 (Protected)
            -> 0.2 for Gender: 2 (Protected)
            -> 0.1 for Gender: 3 (Protected)

    - alpha: Percentage of Type I error

    - color_blind: specifies if this is color blind ranking or not.
      True for generating color blind ranking. For applying fair ranking algorithm with different Protected Groups, False.

"""


# multinomFair(data_set, {'Group':4}, 'Score',500, [0.4,0.3,0.2,0.1], float(0.1), False)
@timer_func
def multinomFair(data_set, attributeNamesAndCategories, attributeQuality, k, p, alpha, color_blind):
    categories = determineGroups(attributeNamesAndCategories)
    num_categories = len(categories)
    count = [0 for _ in range(num_categories)];
    ranking = []
    random_choice = ""
    chosen = 0
    num = 1
    swap = 0
    target = ()
    separateByGroups = [[] for _ in range(num_categories)]
    attributeItems = attributeNamesAndCategories.items()

    no_remaining_items = 0

    #array of the attribute names
    attributes = [i[0] for i in attributeItems]

    #saves the generated ranking in .txt file
    f = open('ranking.txt', 'w')

    # Separate the groups into each list
    separateByGroups = separate_groups(data_set, categories, attributeItems)

    # Sort the items in ascending order of their quality
    for i in range(num_categories):
        separateByGroups[i] = sorted(separateByGroups[i], key=lambda item: item[attributeQuality])

    if(color_blind):
        minimum_targets = numpy.zeros((k,num_categories-1))
    else:
        minimum_targets = get_minimum_targets(categories, p, alpha, k)

    for i in minimum_targets:
        for j in range(len(i)):
            if(i[j] > count[j+1] and chosen == 0):
                if(len(separateByGroups[j+1]) == 0):
                    no_remaining_items = 1
                    break
                item = separateByGroups[j+1].pop()
                ranking.append(item)
                f.write(str(item)+"\n")
                count[j+1] += 1;
                chosen = 1;
#                 print item
                if(any(count[a+1]<i[a] for a in range(len(i)))):
                    swap = len(ranking)-1;
                    # for edge case
                    while(any(count[a+1]<i[a] for a in range(len(i)))):
                        target = find_target(i, count, categories);
                        if(count[0] == 0):
                            achieved_target = find_achieved_target(i, count, categories);
                            if (all(ranking[swap][attributes[i]] == achieved_target[i] for i in range(len(attributes)))):
                                print ("item to swap: ",ranking[swap])
                                separateByGroups[categories.index(achieved_target)].append(ranking[swap])
                                ranking = ranking[:swap] + ranking[swap+1 :]
                                ranking.append(separateByGroups[categories.index(target)].pop())
                                count[categories.index(target)] += 1
                                count[categories.index(achieved_target)] -= 1
                        else:
                            print ("Target group: ",target)
                            if (all(ranking[swap][a] == 0 for a in attributes)):
                                separateByGroups[0].append(ranking[swap])
                                ranking = ranking[:swap] + ranking[swap+1 :]
                                ranking.append(separateByGroups[categories.index(target)].pop())
                                count[categories.index(target)] += 1
                                count[0] -= 1;
                        swap -= 1;
        if(no_remaining_items == 1):
            break
        if(chosen == 0):
            heads = []
            first = 0
            duplicates=[]
            for x in separateByGroups:
                if (len(x) == 0):
                    heads.append(-100)
                else:
                    heads.append(x[len(x)-1][attributeQuality])
            most_val = max(heads)
            # handle cases when more than one group with the best score exists
            for l in range(len(heads)):
                if (heads[l]==most_val):
                    duplicates.append(l)

            if(len(duplicates)>1):
                if(color_blind):
                    idx = numpy.random.choice(duplicates)
                else:
                    p_random = [p[r] for r in duplicates]
                    p_random = numpy.array(p_random)/sum(p_random)
                    idx = numpy.random.choice(duplicates, 1, p=p_random)[0]
            else:
                idx= duplicates[0]
            item = separateByGroups[idx].pop()
#             print item
            f.write(str(item)+"\n")
            ranking.append(item);
            count[idx] += 1;
        print (num, ": ", count, "Minimum Target: ",minimum_targets[num-1]);
        f.write(str(num)+": "+str(count)+"Minimum Target: "+str(minimum_targets[num-1])+"\n")
        print ("CDF: ", multinomCDF(num_categories, num, p, [num]+count[1:]))
        f.write("CDF: "+str(multinomCDF(num_categories, num, p, [num]+count[1:]))+"\n")
        num = num+1;
        chosen = 0;
    f.close()
    rest = []
    best = ranking[0][attributeQuality]
    for i in separateByGroups:
        for j in i:
            rest.append(float(j[attributeQuality])/float(best))
    plot(ranking,attributeNamesAndCategories, attributeQuality)
    plot_scatter(ranking,attributeNamesAndCategories, attributeQuality)

    for q in range(len(ranking)):
        ranking[q]['k']=q+1
        ranking[q]['Utility'] = (1.0/numpy.log(1 + q+1))*ranking[q][attributeQuality]

    # generates a .docx file that shows the generated ranking, in which each group is shown with different color.
    # the colors are set for maximum 6 different groups. If there are more than 6 groups, the colors have to be configured for those groups.
    if(num_categories <=6):
        document = Document()
        document.add_heading("Output Ranking", 0)

        for idx in range(len(ranking)):
            if (all(ranking[idx][attributes[i]] == categories[0][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)

            elif (all(ranking[idx][attributes[i]] == categories[1][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0xcc, 0x00, 0x00)

            elif (all(ranking[idx][attributes[i]] == categories[2][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0x00, 0x66, 0xff)

            elif (all(ranking[idx][attributes[i]] == categories[3][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0x00, 0x99, 0x33)

            elif (all(ranking[idx][attributes[i]] == categories[4][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0xFF, 0xCC, 0x33)

            elif (all(ranking[idx][attributes[i]] == categories[5][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0x99, 0x66, 0x00)

            elif (all(ranking[idx][attributes[i]] == categories[6][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0xff, 0x00, 0xff)

            elif (all(ranking[idx][attributes[i]] == categories[7][i] for i in range(len(attributes)))):
                run = document.add_paragraph().add_run(str(ranking[idx]))
                font = run.font
                font.color.rgb = RGBColor(0xff, 0x00, 0xff)

        document.save('output_ranking_different_color.docx')

    return ranking

def parseDataset(data_set, attributeName, categories):
    json_data = open(data_set).read()
    data = numpy.array(json_data.split("\n"))
    data_set = []
    group_count = numpy.zeros(categories)

    f = open('properties.txt', 'w')

    for i in data:
        data_set.append(ast.literal_eval(i))

    for j in range(len(data_set)):
        data_set[j]['Index']=j+1
        for i in range(categories):
            if(data_set[j][attributeNames]==i):
                group_count[i] = group_count[i]+1

    for i in range(categories):
        f.write("Group "+str(i)+": "+str(group_count[i])+"\n")
        print ("Group ",i,": ",group_count[i])

    f.write("Size of data set: "+str(len(data_set))+"\n")
    print ("Size of data set: ", len(data_set))
    f.close()

    return data_set



def child_generator_4(G, k, p ,a, tau):
    tau_p = [k] + list(tau) ;
    temp = copy.copy(tau_p)
    old_cdf = multinomCDF(G, k, p, tau_p)
    new_cdf = 0;
    initial = 1;
    not_fulfilled = 0;
    diz = {}
    ls = []
    i = 0

    if old_cdf > a:
        return  tau

    for i in range(len(tau_p)-1):
        temp[i+1] = temp[i+1]+1;
        if(initial == 1):
            tau_p = copy.copy(temp);
            cdf = multinomCDF(G, k, p, tau_p);
            initial = 0;
            if (cdf > a):

                ls.append(tau_p[1:])



#             print(f"primo if: {tau_p,temp, cdf}");

        else:
            new_cdf = multinomCDF(G, k, p, temp)
            if(new_cdf > a):
                tau_p = copy.copy(temp);
                cdf = multinomCDF(G, k, p, tau_p)

                ls.append(tau_p[1:])
#                   print(f"secondopr_if: {tau_p,temp,cdf,new_cdf}")


#                 ls.append(tau_p);


#             print(f"secondo if: {tau_p,temp,cdf,new_cdf}")
        if(new_cdf >= a or cdf >= a or cdf <= a or new_cdf< a):
            temp[i+1] = temp[i+1]-1


    return ls

def countElements(group, list_):
    if group == 0:
        lista =  copy.copy(list_)
        return lista

    elif group == 1:

        list_[0] = list_[0] + 1
        lista =  copy.copy(list_)
        return lista


    elif group == 2:
        list_[1] = list_[1] + 1
        lista =  copy.copy(list_)
        return lista


    elif group == 3:
        list_[2] = list_[2] + 1
        lista =  copy.copy(list_)
        return lista




# if __name__ == '__main__'
#     import argparse
#     parser = argparse.ArgumentParser(description="data_set, attributeNamesAndCategories, attributeQuality, k, p, alpha, color_blind")
#     parser.add_argument('--data_set', metavar='path', required=True,
#                        help='path to the data set JSON file')
#     parser.add_argument('--attributeNamesAndCategories', required=True,
#                        help='protected attribute name and size in dictionary')
#     parser.addargument('--attributeQuality', required=True)
#     parser.addargument('--k', required=True)
#     parser.addargument('--p', required=True)
#     parser.addargument('--alpha',required=True)
#     parser.addargument('--color_blind',required=True)
#     args = parser.parse_args()

#     data = parseDataset(args.data_set)

#     multinomFair(data_set=data, attributeNamesAndCategories=args.attributeNamesAndCategories,
#                  attributeQuality=args.attributeQuality, k=args.k, p=args.p,
#                  alpha=args.alpha, color_blind=args.color_blind)



alpha_list = numpy.arange(0.01, 0.21, 0.01)
L_list = numpy.arange(0, 1.01, 0.01)
intervals_utility = []
intervals_alpha = []
alpha = 0.1
attributeQuality = 'Score'
res = []
confidence = 0.90
for L in L_list:

    json_data = open('.\german_credit.json').read()
    data = numpy.array(json_data.split("\n"))
    data_set = []
    ## Clean the data
    for i in data:
        data_set.append(ast.literal_eval(i))
    attribute = {'Group':4}


    #initialization
#     alpha = 0.1
    p = [0.4,0.3,0.2,0.1]
    k_th = 51
    ## Definiamo un parametro di loss L per bilanciare tra fairness e utilità

    attributes = {'Group':4}
    #Sort Values
    data_sert = sorted(data_set, reverse=True,key = lambda user: (user['Score']))
    data_set = data_sert

    for diz, i in zip(data_set, range(len(data_set))):
        diz['k'] = i + 1


    cb_rank = pd.DataFrame(data_set)
    cb_rank['DCG_PreR'] = cb_rank['Score']/(numpy.log(2 + cb_rank['k']))
    new_rank, exposureByGroups, countByGroups, exposureStd, alphaByGroups, DCGByGroups = Greedy_Wise_Utility.GWU(data_set, p, alpha, k_th, attributeQuality, attributes, 'Group', L)

    new_rank['PreR_DCG'] = new_rank['Score']/(numpy.log(2 + new_rank['k']))
    new_rank['PostR_DCG'] = new_rank['Score']/(numpy.log(2 + new_rank['new_k']))
    new_rank['Utility_Loss_individual'] = new_rank['PostR_DCG']-new_rank['PreR_DCG']
    new_rank['Utility_Loss_position'] = new_rank['PostR_DCG']-cb_rank['DCG_PreR']
    new_rank['ULP*ULI'] = new_rank['Utility_Loss_position']*new_rank['Utility_Loss_individual']

    # calculate the kendall's correlation between two variables
    # calculate kendall's correlation
    coef, p = kendalltau(new_rank['PostR_DCG'], cb_rank['DCG_PreR'].head(50))
#     print('Kendall correlation coefficient: %.3f' % coef)
    # interpret the significance
#     alpha = 0.1
#     if p > alpha:
#         print('Samples are uncorrelated (fail to reject H0) p=%.3f' % p)
#     else:
#         print('Samples are correlated (reject H0) p=%.3f' % p)

    diz_results = {}
    diz_exposure = {}
    diz_results['L'] = L
    diz_results['Utility'] = new_rank['Utility_adj'].sum()
    diz_results['Utility_mean'] = new_rank['Utility_adj'].mean()
    diz_results['Utility_std'] = new_rank['Utility_adj'].std()


    diz_results['alpha_adj'] = new_rank['alpha_adj'].mean()
    diz_results['alphaAdjStd'] = new_rank['alpha_adj'].std()
    dof = len(new_rank)-1


    t_crit = numpy.abs(t.ppf((1-confidence)/2,dof))


    diz_exposure['exposureByGroups'] = exposureByGroups
    diz_exposure['L'] = L
    diz_exposure['countByGroups'] = countByGroups
    diz_exposure['exposureStd'] = exposureStd
    diz_results['ExposureByGroups'] = diz_exposure


    interval_utility =(diz_results['Utility_mean']-diz_results['Utility_std']*t_crit/numpy.sqrt(dof+1)
                      ,diz_results['Utility_mean']+diz_results['Utility_std']*t_crit/numpy.sqrt(dof+1))



    interval_alpha = (diz_results['alpha_adj']-diz_results['alphaAdjStd']*t_crit/numpy.sqrt(dof+1)
                     ,diz_results['alpha_adj']+diz_results['alphaAdjStd']*t_crit/numpy.sqrt(dof+1))



#     new_rank.groupby
    diz_results['lower_alpha_int'] = interval_alpha[0]
    diz_results['upper_alpha_int'] = interval_alpha[1]

    intervals_alpha.append(interval_alpha)
    intervals_utility.append(interval_utility)
    diz_results['Kendal_tau'] = coef


    res.append(diz_results)

df_res = pd.DataFrame(res)



exposureData = []
intervalsExposure = []
for row in df_res['ExposureByGroups']:
    count = 0
    for meanExp, countExp, exposureStd in zip(row['exposureByGroups'], row['countByGroups'], row['exposureStd']):
        diz = {}
        diz['Group'] = count
        diz['averageExposure'] = meanExp
        diz['lenGroup'] = countExp -1
        diz['exposureStd'] = exposureStd
        t_crit = numpy.abs(t.ppf((1-confidence)/2,dof))

        interval_utility = (diz['averageExposure'] - diz['exposureStd']*t_crit/numpy.sqrt(diz['lenGroup'])
        ,diz['averageExposure'] + diz['exposureStd'] * t_crit/numpy.sqrt(diz['lenGroup']))

        diz['upper_int']=interval_utility[1]
        diz['lower_int'] = interval_utility[0]

#         intervalsExposure.append(interval_utility)
        diz['L'] = row['L']
#         diz['alpha_t'] = alpha
        count = count + 1
        exposureData.append(diz)

exposureDf = pd.DataFrame(exposureData)


f = plt.figure(figsize=(10, 10))
sns.scatterplot(x=(exposureDf['L']), y=exposureDf['averageExposure'], palette="tab10", hue=exposureDf['Group'])

plt.savefig('./Plot/avgExposure_VS_L.png')




plt.figure(figsize=(8,8))

plt.errorbar(x=df_res['L'],
             y=df_res['alpha_adj'],
             yerr=[(top-bot)/2 for bot,top in intervals_alpha],
             fmt='o')

plt.xlabel('L')
plt.ylabel('alpha_c')


plt.savefig('./Plot/alpha_c_VS_L.png')


f = plt.figure(figsize=(8,8))
sns.scatterplot(x=(df_res['L']), y=(df_res['alpha_adj']), palette="tab10")

plt.savefig('./Plot/alpha_ADJ_VS_L.png')



f = plt.figure(figsize=(8,8))
sns.scatterplot(x=(df_res['L']), y=(df_res['Kendal_tau']), palette="tab10")
plt.savefig('./Plot/KendalTau_VS_L.png')



# Sensitivity analysis for L and alpha_t parameter


alpha_list = numpy.arange(0.01, 0.21, 0.01)
L_list = numpy.arange(0, 1.01, 0.01)
import Greedy_Wise_Utility
from numpy.random import rand
from numpy.random import seed
from scipy.stats import kendalltau
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import t
from mpl_toolkits.mplot3d import Axes3D
from matplotlib.colors import ListedColormap


intervals_utility = []
intervals_alpha = []

attributeQuality = 'Score'
res = []
confidence = 0.90

for alpha in alpha_list:
    for L in L_list:

        json_data = open('.\german_credit.json').read()
        data = numpy.array(json_data.split("\n"))
        data_set = []
        ## Clean the data
        for i in data:
            data_set.append(ast.literal_eval(i))
        attribute = {'Group':4}


        #initialization
#         alpha = 0.1
        p = [0.4,0.3,0.2,0.1]
        k_th = 51
        ## Le'ts define a loss parameter L to chose if prefer much fairness or utility

        attributes = {'Group':4}
        #Sort Values
        data_sert = sorted(data_set, reverse=True,key = lambda user: (user['Score']))
        data_set = data_sert

        for diz, i in zip(data_set, range(len(data_set))):
            diz['k'] = i + 1


        cb_rank = pd.DataFrame(data_set)
        cb_rank['DCG_PreR'] = cb_rank['Score']/(numpy.log(2 + cb_rank['k']))
        new_rank, exposureByGroups, countByGroups, exposureStd, alphaByGroups, DCGByGroups = Greedy_Wise_Utility.GWU(data_set, p, alpha, k_th, attributeQuality, attributes, 'Group', L)
        new_rank['alpha_thresh'] = alpha
        new_rank['PreR_DCG'] = new_rank['Score']/(numpy.log(2 + new_rank['k']))
        new_rank['PostR_DCG'] = new_rank['Score']/(numpy.log(2 + new_rank['new_k']))
        new_rank['Utility_Loss_individual'] = new_rank['PostR_DCG']-new_rank['PreR_DCG']
        new_rank['Utility_Loss_position'] = new_rank['PostR_DCG']-cb_rank['DCG_PreR']
        new_rank['ULP*ULI'] = new_rank['Utility_Loss_position']*new_rank['Utility_Loss_individual']

        # calculate the kendall's correlation between two variables
        # calculate kendall's correlation
        coef, p = kendalltau(new_rank['PostR_DCG'], cb_rank['DCG_PreR'].head(50))
    #     print('Kendall correlation coefficient: %.3f' % coef)
        # interpret the significance
    #     alpha = 0.1
    #     if p > alpha:
    #         print('Samples are uncorrelated (fail to reject H0) p=%.3f' % p)
    #     else:
    #         print('Samples are correlated (reject H0) p=%.3f' % p)

        diz_results = {}
        diz_exposure = {}
        diz_DCG = {}
        diz_alpha = {}
        diz_results['L'] = L
        diz_results['DCG'] =  new_rank['PostR_DCG'].mean()
        diz_results['Utility'] = new_rank['Utility_adj'].sum()
        diz_results['Utility_mean'] = new_rank['Utility_adj'].mean()
        diz_results['Utility_std'] = new_rank['Utility_adj'].std()
        diz_results['alpha'] = alpha
        diz_results['kendal_tau'] = coef
        diz_results['alpha_adj'] = new_rank['alpha_adj'].mean()
        diz_results['alphaAdjStd'] = new_rank['alpha_adj'].std()
        dof = len(new_rank)-1


        t_crit = numpy.abs(t.ppf((1-confidence)/2,dof))

        diz_alpha['alphaByGroups'] = alphaByGroups
        diz_DCG['DCGByGroups'] = DCGByGroups
        diz_exposure['exposureByGroups'] = exposureByGroups
        diz_exposure['L'] = L
        diz_exposure['alpha'] = alpha
        diz_exposure['countByGroups'] = countByGroups
        diz_exposure['exposureStd'] = exposureStd
        diz_results['alphaByGroups'] = diz_alpha
        diz_results['ExposureByGroups'] = diz_exposure
        diz_results['DCGByGroups'] = diz_DCG

        interval_utility =(diz_results['Utility_mean']-diz_results['Utility_std']*t_crit/numpy.sqrt(dof+1)
                          ,diz_results['Utility_mean']+diz_results['Utility_std']*t_crit/numpy.sqrt(dof+1))



        interval_alpha = (diz_results['alpha_adj']-diz_results['alphaAdjStd']*t_crit/numpy.sqrt(dof+1)
                         ,diz_results['alpha_adj']+diz_results['alphaAdjStd']*t_crit/numpy.sqrt(dof+1))



    #     new_rank.groupby
        diz_results['lower_alpha_int'] = interval_alpha[0]
        diz_results['upper_alpha_int'] = interval_alpha[1]

        intervals_alpha.append(interval_alpha)
        intervals_utility.append(interval_utility)
        diz_results['Kendal_tau'] = coef


        res.append(diz_results)

df_res = pd.DataFrame(res)

df_res['ExposureByGroups'][0]['exposureByGroups'][0]
df_res['DTR_1'] = df_res['ExposureByGroups'].apply(lambda x: x['exposureByGroups'][0]/x['exposureByGroups'][1] - 1)
df_res['DTR_2'] = df_res['ExposureByGroups'].apply(lambda x: x['exposureByGroups'][0]/x['exposureByGroups'][2] - 1)
df_res['DTR_3'] = df_res['ExposureByGroups'].apply(lambda x: x['exposureByGroups'][0]/x['exposureByGroups'][3] - 1)
df_res['Metric'] = df_res['kendal_tau']*df_res['alpha_adj'] - (df_res['DTR_1']*0.3 + df_res['DTR_2']*0.2 + df_res['DTR_3']*0.1)
df_res.loc[df_res['Metric'].idxmax()]
exposureData = []
intervalsExposure = []
for row, row1, row2 in zip(df_res['ExposureByGroups'], df_res['alphaByGroups'], df_res['DCGByGroups']):
    count = 0
    for meanExp, countExp, exposureStd, meanAlpha, meanDCG in zip(row['exposureByGroups'], row['countByGroups'], row['exposureStd'], row1['alphaByGroups'], row2['DCGByGroups']):
        diz = {}
        diz['Group'] = count
        diz['averageExposure'] = meanExp
        diz['alpha_c'] = meanAlpha
        diz['meanDCG'] = meanDCG
        diz['lenGroup'] = countExp -1
        diz['exposureStd'] = exposureStd
        t_crit = numpy.abs(t.ppf((1-confidence)/2,dof))

        interval_utility =(diz['averageExposure']-diz['exposureStd']*t_crit/numpy.sqrt(diz['lenGroup'])
        ,diz['averageExposure']+diz['exposureStd']*t_crit/numpy.sqrt(diz['lenGroup']))

        diz['upper_int']=interval_utility[1]
        diz['lower_int'] = interval_utility[0]

#         intervalsExposure.append(interval_utility)
        diz['L'] = row[ 'L']
        diz['alpha'] = row['alpha']
        count = count + 1
        exposureData.append(diz)

exposureDf = pd.DataFrame(exposureData)

# exposureDf.to_csv('exposureDf.csv')
exposureDf.to_excel('exposureDf.xlsx')




import Greedy_Wise_Utility
from numpy.random import rand
from numpy.random import seed
from scipy.stats import kendalltau
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import t
from mpl_toolkits.mplot3d import Axes3D
from matplotlib.colors import ListedColormap
fig = plt.figure(figsize=(20,20))
ax = fig.add_subplot(projection='3d')

ax.set_title("Exposure at variation of alpha and L",
             fontsize=20)
ax.set_xlabel("L",
              fontsize=16)
ax.set_ylabel("alpha",
              fontsize=16)
ax.set_zlabel("Exposure",
              fontsize=16)


cmap = ListedColormap(sns.color_palette("tab10").as_hex())
sc = ax.scatter( exposureDf['L'], exposureDf['alpha'],exposureDf['averageExposure'], c=exposureDf['Group'], cmap = cmap, alpha=1)
# handles, labels = ax.get_legend_handles_labels()

plt.legend(*sc.legend_elements(), bbox_to_anchor = (0, 1), loc=2)


ax.text(2.1,9.4,5.0,
       "1",fontsize= 10)

ax.view_init(elev=15., azim=45)


plt.savefig('./Plot/Exposure_at_L_and_alpha_variation')
# er i gruppi protetti inserire un colore in 'crescendo da chiaro a scuro'
